#!/usr/bin/env python3
"""
Word Document Fill API
API để chèn dữ liệu vào file Word sử dụng doc2docx
"""

from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import tempfile
from datetime import datetime
import logging
from docx import Document
import re
import subprocess

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Enable CORS for all routes
CORS(app)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class WordFiller:
    def __init__(self):
        # Chỉ tập trung vào 7 trường chính theo yêu cầu
        self.field_mappings = {
            # Số CCCD
            "Số CCCD": "so_cccd",
            "CCCD": "so_cccd",
            "Căn cước công dân": "so_cccd",
            "Số căn cước": "so_cccd",
            "Số căn cước công dân": "so_cccd",
            "Số CMND hoặc căn cước công dân": "so_cccd",


            # Số CMND
            "Số CMND": "so_cmnd",
            "CMND": "so_cmnd",
            "Chứng minh nhân dân": "so_cmnd",
            "Số chứng minh": "so_cmnd",
            "Số chứng minh nhân dân": "so_cmnd",

            # Họ và tên
            "Họ và tên": "ho_ten",
            "Họ, chữ đệm, tên": "ho_ten",
            "Họ tên": "ho_ten",
            "Họ, chữ đệm, tên người yêu cầu": "ho_ten",
            "Tên": "ho_ten",

            # Giới tính
            "Giới tính": "gioi_tinh",
            "Phái": "gioi_tinh",
            "Nam/Nữ": "gioi_tinh",

            # Ngày sinh
            "Ngày sinh": "ngay_sinh",
            "Ngày, tháng, năm sinh": "ngay_sinh",
            "Sinh ngày": "ngay_sinh",
            "Năm sinh": "ngay_sinh",

            # Nơi cư trú
            "Nơi cư trú": "noi_cu_tru",
            "Địa chỉ cư trú": "noi_cu_tru",
            "Chỗ ở hiện tại": "noi_cu_tru",
            "Địa chỉ": "noi_cu_tru",

            # Ngày cấp CCCD
            "Ngày cấp CCCD": "ngay_cap_cccd",
            "Ngày cấp": "ngay_cap_cccd",
            "Cấp ngày": "ngay_cap_cccd",
            "Ngày cấp căn cước": "ngay_cap_cccd"
        }

    def convert_doc_to_docx(self, doc_path):
        """Chuyển đổi .doc sang .docx giữ nguyên định dạng"""
        try:
            print(f"🔄 Converting {os.path.basename(doc_path)} to .docx...")

            # Đảm bảo chỉ xử lý file .doc
            if not doc_path.lower().endswith('.doc'):
                print(f"❌ File không phải .doc: {doc_path}")
                return None

            docx_path = doc_path.replace('.doc', '_converted.docx')

            # Method 1: LibreOffice headless (giữ được định dạng tốt nhất)
            if self._try_libreoffice_headless(doc_path, docx_path):
                return docx_path

            # Method 2: Pandoc (giữ được một số định dạng)
            if self._try_pandoc_conversion(doc_path, docx_path):
                return docx_path

            # Method 3: unoconv (LibreOffice API)
            if self._try_unoconv_conversion(doc_path, docx_path):
                return docx_path

            print("❌ All conversion methods failed")
            print("💡 Để giữ định dạng tốt nhất, hãy cài LibreOffice:")
            print("   - macOS: brew install --cask libreoffice")
            print("   - Ubuntu: sudo apt install libreoffice")
            print("   - Windows: Download từ libreoffice.org")
            return None

        except Exception as e:
            print(f"❌ Conversion error: {e}")
            return None

    def _try_libreoffice_headless(self, doc_path, docx_path):
        """LibreOffice headless - giữ nguyên định dạng tốt nhất"""
        try:
            print("  🔧 Trying LibreOffice headless...")

            # Tìm LibreOffice executable
            libreoffice_paths = [
                'libreoffice',  # Nếu có trong PATH
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
                '/usr/bin/libreoffice',  # Linux
                '/opt/libreoffice/program/soffice',  # Linux alternative
                'soffice'  # Alternative command
            ]

            libreoffice_cmd = None
            for path in libreoffice_paths:
                try:
                    result = subprocess.run([path, '--version'],
                                            capture_output=True, timeout=5)
                    if result.returncode == 0:
                        libreoffice_cmd = path
                        print(f"    Found LibreOffice at: {path}")
                        break
                except:
                    continue

            if not libreoffice_cmd:
                print("    ❌ LibreOffice executable not found")
                return False

            # Sử dụng các tham số tối ưu để giữ định dạng
            result = subprocess.run([
                libreoffice_cmd,
                '--headless',           # Không mở GUI
                '--invisible',          # Chạy ẩn
                '--nodefault',         # Không load default document
                '--nolockcheck',       # Không check file lock
                '--convert-to', 'docx', # Convert sang docx
                '--outdir', os.path.dirname(doc_path),  # Output directory
                doc_path
            ], capture_output=True, text=True, timeout=45)

            if result.returncode == 0:
                # LibreOffice tạo file với tên gốc + .docx
                auto_path = doc_path.replace('.doc', '.docx')
                if os.path.exists(auto_path):
                    # Rename để match expected output path
                    if auto_path != docx_path:
                        os.rename(auto_path, docx_path)
                    print(f"  ✅ LibreOffice conversion successful (format preserved)")
                    return True

            print(f"  ❌ LibreOffice failed: {result.stderr}")
            return False

        except subprocess.TimeoutExpired:
            print(f"  ❌ LibreOffice timeout (file quá lớn hoặc phức tạp)")
            return False
        except FileNotFoundError:
            print(f"  ❌ LibreOffice not installed")
            return False
        except Exception as e:
            print(f"  ❌ LibreOffice error: {e}")
            return False

    def _try_pandoc_conversion(self, doc_path, docx_path):
        """Pandoc conversion - giữ được một số định dạng"""
        try:
            print("  🔧 Trying Pandoc...")

            result = subprocess.run([
                'pandoc',
                doc_path,
                '-o', docx_path,
                '--from=doc',
                '--to=docx'
            ], capture_output=True, text=True, timeout=30)

            if result.returncode == 0 and os.path.exists(docx_path):
                print(f"  ✅ Pandoc conversion successful")
                return True

            print(f"  ❌ Pandoc failed: {result.stderr}")
            return False

        except subprocess.TimeoutExpired:
            print(f"  ❌ Pandoc timeout")
            return False
        except FileNotFoundError:
            print(f"  ❌ Pandoc not installed")
            return False
        except Exception as e:
            print(f"  ❌ Pandoc error: {e}")
            return False

    def _try_unoconv_conversion(self, doc_path, docx_path):
        """unoconv - LibreOffice API"""
        try:
            print("  🔧 Trying unoconv...")

            result = subprocess.run([
                'unoconv',
                '-f', 'docx',
                '-o', docx_path,
                doc_path
            ], capture_output=True, text=True, timeout=30)

            if result.returncode == 0 and os.path.exists(docx_path):
                print(f"  ✅ unoconv conversion successful")
                return True

            print(f"  ❌ unoconv failed: {result.stderr}")
            return False

        except subprocess.TimeoutExpired:
            print(f"  ❌ unoconv timeout")
            return False
        except FileNotFoundError:
            print(f"  ❌ unoconv not installed")
            return False
        except Exception as e:
            print(f"  ❌ unoconv error: {e}")
            return False



    def process_text_line(self, text, data_dict):
        """Xử lý từng dòng text để thay thế dữ liệu"""
        if not text.strip():
            return text

        new_text = text

        for field_name, data_key in self.field_mappings.items():
            if field_name in text and data_key in data_dict:

                # Pattern 1: "Tên trường: giá_trị_cũ(số)" -> "Tên trường: giá_trị_mới"
                pattern1 = f"({re.escape(field_name)}):\\s*[^\\(\\n]*\\([^\\)]*\\)"
                if re.search(pattern1, text):
                    new_text = re.sub(pattern1, f"\\1: {data_dict[data_key]}", new_text)
                    return new_text

                # Pattern 2: "Tên trường: giá_trị_cũ" -> "Tên trường: giá_trị_mới"
                pattern2 = f"({re.escape(field_name)}):\\s*[^\\n]+"
                if re.search(pattern2, text) and not text.endswith(':'):
                    new_text = re.sub(pattern2, f"\\1: {data_dict[data_key]}", new_text)
                    return new_text

                # Pattern 3: "Tên trường: " (trống) -> "Tên trường: giá_trị_mới"
                pattern3 = f"({re.escape(field_name)}):\\s*$"
                if re.search(pattern3, text):
                    new_text = re.sub(pattern3, f"\\1: {data_dict[data_key]}", new_text)
                    return new_text

        return new_text

    def fill_document(self, docx_path, data_dict):
        """Điền dữ liệu vào document"""
        print(f"📄 Processing: {os.path.basename(docx_path)}")

        doc = Document(docx_path)
        replacements_made = 0

        # Xử lý paragraphs
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            new_text = self.process_text_line(original_text, data_dict)
            if new_text != original_text:
                paragraph.text = new_text
                replacements_made += 1
                print(f"  ✅ Updated: {original_text[:40]}... → {new_text[:40]}...")

        # Xử lý tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    new_text = self.process_text_line(original_text, data_dict)
                    if new_text != original_text:
                        cell.text = new_text
                        replacements_made += 1
                        print(f"  ✅ Updated table: {original_text[:40]}... → {new_text[:40]}...")

        # Lưu document đã điền dữ liệu
        # Tạo tên file output an toàn
        if docx_path.endswith('.docx'):
            output_path = docx_path[:-5] + '_filled.docx'  # Bỏ .docx cuối và thêm _filled.docx
        else:
            output_path = docx_path + '_filled.docx'
        doc.save(output_path)

        print(f"\n🎉 Success!")
        print(f"📊 Made {replacements_made} replacements")
        print(f"💾 Output: {os.path.basename(output_path)}")

        return output_path, replacements_made

# Khởi tạo word filler
filler = WordFiller()

@app.route('/')
def index():
    """Trang chủ - Web interface"""
    return render_template('index.html')

@app.route('/health', methods=['GET'])
def health_check():
    """Kiểm tra trạng thái API"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'service': 'Word Document Fill API',
        'version': '1.0.0'
    })

@app.route('/fill', methods=['POST'])
def fill_document():
    """API endpoint để điền dữ liệu vào file Word"""
    
    # Kiểm tra file upload
    if 'file' not in request.files:
        return jsonify({'error': 'Không có file được upload'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Không có file được chọn'}), 400
    
    # Kiểm tra định dạng file
    if not (file.filename.lower().endswith('.doc') or file.filename.lower().endswith('.docx')):
        return jsonify({'error': 'Chỉ hỗ trợ file .doc và .docx'}), 400
    
    # Lấy dữ liệu từ request
    data = {}
    if request.is_json:
        data = request.get_json()
    else:
        data = request.form.to_dict()
    
    if not data:
        return jsonify({'error': 'Không có dữ liệu để điền'}), 400
    
    try:
        # Tạo thư mục tạm
        temp_dir = tempfile.mkdtemp()
        filename = secure_filename(file.filename)
        
        if filename.lower().endswith('.doc'):
            # Lưu file .doc và chuyển đổi sang .docx
            doc_path = os.path.join(temp_dir, filename)
            file.save(doc_path)
            
            docx_path = filler.convert_doc_to_docx(doc_path)
            if not docx_path:
                return jsonify({'error': 'Chuyển đổi file .doc thất bại'}), 500
        
        elif filename.lower().endswith('.docx'):
            # Sử dụng file .docx trực tiếp
            docx_path = os.path.join(temp_dir, filename)
            file.save(docx_path)
        
        # Điền dữ liệu vào document
        output_path, replacements = filler.fill_document(docx_path, data)
        
        # Trả về file đã điền dữ liệu
        # Tạo tên file output đúng định dạng
        if filename.lower().endswith('.doc'):
            output_filename = f"filled_{filename.replace('.doc', '.docx')}"
        elif filename.lower().endswith('.docx'):
         
            # Xử lý file .docx: bỏ .docx và thêm filled_.docx
            base_name = filename[:-5]  # Bỏ .docx cuối
            output_filename = f"filled_{base_name}.docx"
            
        else:
            output_filename = f"filled_{filename}"
        print(output_filename)
        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_filename
        )
    
    except Exception as e:
        logger.error(f"Fill error: {e}")
        return jsonify({'error': f'Lỗi xử lý file: {str(e)}'}), 500

@app.route('/convert', methods=['POST'])
def convert_only():
    """API endpoint chỉ để chuyển đổi .doc sang .docx"""
    
    if 'file' not in request.files:
        return jsonify({'error': 'Không có file được upload'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Không có file được chọn'}), 400
    
    if not file.filename.lower().endswith('.doc'):
        return jsonify({'error': 'Chỉ hỗ trợ file .doc'}), 400
    
    try:
        # Tạo thư mục tạm
        temp_dir = tempfile.mkdtemp()
        filename = secure_filename(file.filename)
        doc_path = os.path.join(temp_dir, filename)
        file.save(doc_path)
        
        # Chuyển đổi sang .docx
        docx_path = filler.convert_doc_to_docx(doc_path)
        if not docx_path:
            return jsonify({'error': 'Chuyển đổi thất bại'}), 500
        
        # Trả về file .docx
        return send_file(
            docx_path, 
            as_attachment=True,
            download_name=filename.replace('.doc', '.docx')
        )
    
    except Exception as e:
        logger.error(f"Convert error: {e}")
        return jsonify({'error': f'Lỗi chuyển đổi: {str(e)}'}), 500

if __name__ == '__main__':
    print("🚀 Word Document Fill API")
    print("=" * 50)
    print("Endpoints:")
    print("  GET  /health   - Kiểm tra trạng thái API")
    print("  POST /convert  - Chuyển đổi .doc sang .docx")
    print("  POST /fill     - Điền dữ liệu vào file Word")
    print("=" * 50)
    print("7 trường dữ liệu hỗ trợ:")
    print("  so_cccd: 123456789012")
    print("  so_cmnd: 123456789")
    print("  ho_ten: Nguyễn Văn A")
    print("  gioi_tinh: Nam/Nữ")
    print("  ngay_sinh: 01/01/1990")
    print("  noi_cu_tru: 123 Đường ABC, Quận 1, TP.HCM")
    print("  ngay_cap_cccd: 15/05/2020")
    print("=" * 50)
    app.run(debug=True, host='0.0.0.0', port=5003)
